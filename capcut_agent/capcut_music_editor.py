"""
CapCut Music Compilation Editor
================================
pycapcut을 사용하여 여러 노래를 하나의 CapCut 프로젝트로 자동 편집합니다.

기능:
- 음악 폴더의 오디오 파일을 순서대로 배치
- 각 노래마다 2초 fade out
- 각 노래 사이 2초 간격 유지
- 배경 폴더에 이미지 또는 영상 파일이 있으면 슬라이드쇼/영상으로 배경 추가 (이미지+영상 혼용 가능)
- 영상이 여러 개인 경우 일정 시간마다 번갈아 표시하며, 전환 시 fade out 2초 적용
- 파일명이 수자/대문자 prefix(AAA, 01...)로 시작하는 노래는 순서대로 앞에, 나머지는 랜덤 셔플
- 각 노래 시작 시 제목 텍스트 표시
- 노래 목록 .docx 문서 자동 생성

사용법:
1. 아래 CONFIG 섹션에서 경로를 설정하세요
2. pip install -r requirements.txt 실행
3. python capcut_music_editor.py 실행
4. CapCut에서 새로 생성된 초안을 열어서 확인
"""

import os
import random
import re
import sys
from pathlib import Path

# ─────────────────────────────────────────────
#  CONFIG — 로컬 환경 변수 (.env) 설정 로드
# ─────────────────────────────────────────────

def load_local_env():
    """.env 파일에서 환경변수를 로드합니다. (python-dotenv 패키지가 없으면 직접 파싱)"""
    env_path = Path(__file__).parent / ".env"
    try:
        from dotenv import load_dotenv
        load_dotenv(env_path)
    except ImportError:
        if env_path.exists():
            with open(env_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#"):
                        continue
                    if "=" in line:
                        key, val = line.split("=", 1)
                        key = key.strip()
                        val = val.strip()
                        # 따옴표 제거
                        if (val.startswith('"') and val.endswith('"')) or (val.startswith("'") and val.endswith("'")):
                            val = val[1:-1]
                        os.environ[key] = val

# 환경변수 로드 실행
load_local_env()

# 음악 파일이 있는 폴더 경로
MUSIC_FOLDER = os.getenv("MUSIC_FOLDER", r"C:\Users\Sung\YouTube-Playlist\Contents\0604")

# 배경 미디어 폴더 경로 (이미지 + 영상 모두 사용 가능, 없으면 None으로 설정)
# BACKGROUND_FOLDER, IMAGES_FOLDER 모두 지원
_bg_env = os.getenv("BACKGROUND_FOLDER") or os.getenv("IMAGES_FOLDER", r"C:\Users\Sung\YouTube-Playlist\Contents\0604")
IMAGES_FOLDER = None if _bg_env in (None, "", "None", "null") else _bg_env

# CapCut 초안 폴더 경로
# CAPCUT_DRAFT_FOLDER 및 CAPCUT_DRAFTS_FOLDER 모두 지원
CAPCUT_DRAFTS_FOLDER = os.getenv("CAPCUT_DRAFT_FOLDER") or os.getenv("CAPCUT_DRAFTS_FOLDER") or r"C:\Users\Sung\AppData\Local\CapCut\User Data\Projects\com.lveditor.draft"

# 생성할 초안 이름
DRAFT_NAME = os.getenv("DRAFT_NAME", "Girl Smile4")

# 영상 해상도
VIDEO_WIDTH  = 1920
VIDEO_HEIGHT = 1080

# 타이밍 설정 (초 단위)
FADE_OUT_SECONDS       = 2    # 각 노래의 오디오 페이드 아웃 길이
GAP_SECONDS            = 2    # 노래 사이 간격
TITLE_SHOW_SECONDS     = 5    # 제목 표시 시간 (노래 길이보다 길면 노래 길이로 자동 조정)
BACKGROUND_CLIP_SECONDS = 30  # 이미지 슬라이드쇼 사용 시 각 이미지를 몇 초마다 전환할지 (영상은 실제 길이 사용)
BACKGROUND_FADE_SECONDS = 2   # 배경 영상 전환 시 fade out 길이 (초)

# 오디오 파일 확장자
AUDIO_EXTENSIONS = {'.mp3', '.wav', '.flac', '.aac', '.m4a', '.ogg', '.opus'}

# 이미지 파일 확장자
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.webp'}

# 영상 파일 확장자
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.avi', '.mkv', '.wmv', '.webm', '.m4v'}

# ─────────────────────────────────────────────
#  의존성 확인
# ─────────────────────────────────────────────

def check_dependencies():
    missing = []
    try:
        import pycapcut
    except ImportError:
        missing.append("pycapcut")
    try:
        import mutagen
    except ImportError:
        missing.append("mutagen")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")

    if missing:
        print("❌ 아래 패키지가 설치되지 않았습니다:")
        for m in missing:
            print(f"   pip install {m}")
        sys.exit(1)

# ─────────────────────────────────────────────
#  오디오 파일 스캔 및 메타데이터 읽기
# ─────────────────────────────────────────────

def get_songs(music_folder: str) -> list[dict]:
    """음악 폴더에서 오디오 파일을 스캔하고 제목/길이 정보를 반환합니다."""
    from mutagen import File as MutagenFile
    from mutagen.mp3 import MP3
    from mutagen.flac import FLAC

    folder = Path(music_folder)
    if not folder.exists():
        print(f"❌ 음악 폴더를 찾을 수 없습니다: {music_folder}")
        sys.exit(1)

    audio_files = sorted(
        [f for f in folder.iterdir() if f.suffix.lower() in AUDIO_EXTENSIONS],
        key=lambda f: f.name
    )

    if not audio_files:
        print(f"❌ 음악 폴더에 오디오 파일이 없습니다: {music_folder}")
        sys.exit(1)

    print(f"\n📂 음악 폴더: {music_folder}")
    print(f"🎵 발견된 파일: {len(audio_files)}개\n")

    songs = []
    for audio_file in audio_files:
        try:
            meta = MutagenFile(str(audio_file), easy=True)
            if meta is None:
                print(f"  ⚠️  건너뜀 (읽기 실패): {audio_file.name}")
                continue

            duration_secs = meta.info.length  # float, in seconds

            # 태그에서 제목 읽기 (없으면 파일명 사용)
            title = audio_file.stem
            if hasattr(meta, 'tags') and meta.tags:
                tag_title = meta.tags.get('title') or meta.tags.get('TIT2')
                if tag_title:
                    title = str(tag_title[0]) if isinstance(tag_title, list) else str(tag_title)

            # "(Remastered)" 또는 "(Remastered 2011)" 등 괄호 포함 패턴 제거
            title = re.sub(r'\s*\(Remastered[^)]*\)', '', title, flags=re.IGNORECASE).strip()

            songs.append({
                'path': str(audio_file),
                'filename': audio_file.name,
                'title': title,
                'duration_secs': duration_secs,
            })
            mins = int(duration_secs // 60)
            secs = int(duration_secs % 60)
            print(f"  ✅ {audio_file.name}  →  {title}  [{mins}:{secs:02d}]")

        except Exception as e:
            print(f"  ⚠️  건너뜀 ({e}): {audio_file.name}")

    print(f"\n총 {len(songs)}개 노래 처리 예정\n")
    return songs


# ─────────────────────────────────────────────
#  노래 순서 정렬 (번호 있으면 앞에, 없으면 셔플)
# ─────────────────────────────────────────────

# 파일명이 이 패턴으로 시작하면 "번호 있음"으로 판단:
# - 3자 이상 대문자 prefix (AAA, AAB, ABA ...)
# - 숫자 prefix (01, 02, 1. ...)
_NUMBERED_PATTERN = re.compile(r'^(?:[A-Z]{3,}|\d+[.\-_ ]?)', re.ASCII)

def sort_and_shuffle_songs(songs: list[dict]) -> list[dict]:
    """제목 번호(prefix) 유무에 따라 노래 순서를 정렬합니다.

    - prefix(AAA/01...)가 있는 노래: 파일명 순 정렬 후 앞에 배치
    - prefix 없는 노래: 랜덤 셔플 후 뒤에 배치
    """
    numbered   = [s for s in songs if _NUMBERED_PATTERN.match(Path(s['filename']).stem)]
    unnumbered = [s for s in songs if not _NUMBERED_PATTERN.match(Path(s['filename']).stem)]

    # 숫자 노래: 파일명 오름차순 정렬 (실제 저장 순서 유지)
    numbered.sort(key=lambda s: s['filename'])

    # 없는 노래: 랜덤 셔플
    random.shuffle(unnumbered)

    result = numbered + unnumbered

    print(f"🔢 순서: 번호 있는 노래 {len(numbered)}개 (파일명 순) + "
          f"셔플된 노래 {len(unnumbered)}개\n")
    return result


# ─────────────────────────────────────────────
#  배경 미디어 파일 스캔 (이미지 + 영상)
# ─────────────────────────────────────────────

ALL_MEDIA_EXTENSIONS = IMAGE_EXTENSIONS | VIDEO_EXTENSIONS

def get_background_media(images_folder) -> list[dict]:
    """배경 폴더에서 이미지 및 영상 파일 목록을 반환합니다.
    
    Returns:
        list of dicts: [{'path': str, 'is_video': bool, 'duration_secs': float|None}, ...]
        duration_secs: 영상이면 실제 길이(초), 이미지면 None
    """
    from mutagen import File as MutagenFile

    if images_folder is None:
        return []
    folder = Path(images_folder)
    if not folder.exists():
        print(f"⚠️  배경 폴더를 찾을 수 없습니다 (건너뜀): {images_folder}")
        return []
    
    media_files = sorted(
        [f for f in folder.iterdir() if f.suffix.lower() in ALL_MEDIA_EXTENSIONS],
        key=lambda f: f.name
    )
    
    result = []
    img_count = 0
    vid_count = 0
    for f in media_files:
        is_video = f.suffix.lower() in VIDEO_EXTENSIONS
        duration_secs = None
        if is_video:
            try:
                meta = MutagenFile(str(f))
                if meta and hasattr(meta, 'info') and hasattr(meta.info, 'length'):
                    duration_secs = meta.info.length
            except Exception:
                pass  # 읽기 실패 시 None 유지
        result.append({'path': str(f), 'is_video': is_video, 'duration_secs': duration_secs})
        if is_video:
            vid_count += 1
            dur_str = f"{duration_secs:.1f}초" if duration_secs else "?초"
            print(f"   🎬 {f.name}  [{dur_str}]")
        else:
            img_count += 1
    
    print(f"🖼️  배경 이미지: {img_count}개 | 🎬 배경 영상: {vid_count}개 (총 {len(result)}개)\n")
    return result


# ─────────────────────────────────────────────
#  CapCut 초안 생성
# ─────────────────────────────────────────────

def build_capcut_draft(songs: list[dict], images: list[str]):
    """pycapcut을 사용하여 CapCut 초안을 생성합니다."""
    import pycapcut as cc
    from pycapcut import trange, tim, SEC

    # 경로 확인
    drafts_path = Path(CAPCUT_DRAFTS_FOLDER)
    if not drafts_path.exists():
        print(f"❌ CapCut 초안 폴더를 찾을 수 없습니다: {CAPCUT_DRAFTS_FOLDER}")
        print("   CapCut 설정 > 초안 위치에서 경로를 확인하세요.")
        sys.exit(1)

    print(f"📁 CapCut 초안 폴더: {CAPCUT_DRAFTS_FOLDER}")
    print(f"📝 초안 이름: {DRAFT_NAME}\n")

    draft_folder = cc.DraftFolder(CAPCUT_DRAFTS_FOLDER)
    try:
        script = draft_folder.create_draft(DRAFT_NAME, VIDEO_WIDTH, VIDEO_HEIGHT, allow_replace=True)
    except PermissionError:
        print(f"❌ '{DRAFT_NAME}' 초안 폴더를 덮어쓸 수 없습니다.")
        print("   CapCut 프로그램이 켜져 있다면 종료하거나, CapCut 홈 화면(프로젝트 목록)으로 이동해 해당 프로젝트 사용을 중단한 뒤 다시 실행해 주세요.")
        sys.exit(1)

    # 트랙 생성 (배경 → 오디오 → 텍스트 순서)
    script.add_track(cc.TrackType.video, "background")
    script.add_track(cc.TrackType.audio, "music")
    script.add_track(cc.TrackType.text,  "titles")

    # 타이밍 값 (마이크로초)
    fade_out_us      = int(FADE_OUT_SECONDS       * SEC)
    gap_us           = int(GAP_SECONDS             * SEC)
    title_max_us     = int(TITLE_SHOW_SECONDS      * SEC)
    bg_clip_us       = int(BACKGROUND_CLIP_SECONDS * SEC)
    bg_fade_us       = int(BACKGROUND_FADE_SECONDS * SEC)

    # 영상 미디어만 따로 추출 (번갈아 사용하기 위해)
    video_media = [m for m in images if m['is_video']]
    has_multi_video = len(video_media) > 1

    current_us = 0  # 현재 타임라인 위치 (마이크로초)

    for i, song in enumerate(songs):
        duration_us = int(song['duration_secs'] * SEC)

        print(f"  [{i+1:3d}/{len(songs)}] {song['title']}")

        # ── 오디오 트랙 ──────────────────────────────
        try:
            audio_seg = cc.AudioSegment(
                song['path'],
                cc.Timerange(current_us, duration_us),
            )
            audio_seg.add_fade(0, fade_out_us)
            script.add_segment(audio_seg, "music")
        except Exception as e:
            print(f"         ⚠️  오디오 추가 실패: {e}")

        # ── 배경 미디어 트랙 (이미지 또는 영상) ────────
        if images:
            total_bg_us = duration_us + gap_us  # 배경이 덮어야 할 전체 구간

            if has_multi_video:
                # ── 영상 여러 개: BACKGROUND_CLIP_SECONDS마다 번갈아 배치 + fade out ──
                seg_start_us = current_us
                vid_idx = i  # 첫 번째 영상 인덱스는 노래 순번 기준으로 시작
                remaining_us = total_bg_us

                while remaining_us > 0:
                    # 해당 영상의 실제 길이를 클립 단위로 사용 (듁록 말마)
                    media = video_media[vid_idx % len(video_media)]
                    vid_duration_secs = media.get('duration_secs')
                    if vid_duration_secs and vid_duration_secs > 0:
                        # mutagen duration과 pycapcut 내부 길이 간 미세 오차 보정 (-200ms 마진)
                        clip_unit_us = max(1, int(vid_duration_secs * SEC) - 200_000)
                    else:
                        clip_unit_us = bg_clip_us  # 길이 정보 없으면 기본값 사용
                    clip_dur_us = min(clip_unit_us, remaining_us)
                    is_last_clip = (remaining_us <= clip_unit_us)

                    media_label = "영상" if media['is_video'] else "이미지"
                    try:
                        bg_seg = cc.VideoSegment(
                            media['path'],
                            cc.Timerange(seg_start_us, clip_dur_us),
                        )
                        # 주의: pycapcut VideoSegment는 add_fade를 지원하지 않음
                        # 영상 전환은 클립이 연속 배치되는 방식으로 처리됨
                        script.add_segment(bg_seg, "background")
                    except Exception as e:
                        print(f"         ⚠️  배경 {media_label} 추가 실패: {e}")

                    seg_start_us += clip_dur_us
                    remaining_us -= clip_dur_us
                    vid_idx += 1

            else:
                # ── 영상/이미지 1개 or 이미지만: 노래 전체 구간에 단일 배치 ──
                media = images[i % len(images)]
                media_path = media['path']
                media_label = "영상" if media['is_video'] else "이미지"
                try:
                    bg_seg = cc.VideoSegment(
                        media_path,
                        cc.Timerange(current_us, total_bg_us),
                    )
                    script.add_segment(bg_seg, "background")
                except Exception as e:
                    print(f"         ⚠️  배경 {media_label} 추가 실패: {e}")

        # ── 제목 텍스트 트랙 ──────────────────────────
        title_duration_us = min(duration_us, title_max_us)
        try:
            text_seg = cc.TextSegment(
                song['title'],
                cc.Timerange(current_us, title_duration_us),
                style=cc.TextStyle(
                    size=8.0,
                    color=(1.0, 1.0, 1.0),   # 흰색
                    bold=True,
                    align=1,                   # 가운데 정렬
                ),
                clip_settings=cc.ClipSettings(transform_y=-0.75),  # 화면 하단
            )
            script.add_segment(text_seg, "titles")
        except Exception as e:
            print(f"         ⚠️  텍스트 추가 실패: {e}")

        # 다음 노래 시작 위치 = 현재 노래 끝 + 간격
        current_us += duration_us + gap_us

    print(f"\n💾 초안 저장 중...")
    script.save()
    print(f"✅ CapCut 초안 생성 완료: '{DRAFT_NAME}'")
    print("   CapCut을 열거나 재시작하여 초안 목록을 새로고침하세요.\n")


# ─────────────────────────────────────────────
#  노래 목록 문서 생성 (.docx)
# ─────────────────────────────────────────────

def generate_song_list_docx(songs: list[dict], output_path: str):
    """노래 제목과 길이가 담긴 Word 문서를 생성합니다."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 제목
    heading = doc.add_heading("🎵 노래 목록 / Song List", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 요약 정보
    total_secs = sum(s['duration_secs'] for s in songs)
    gap_total = GAP_SECONDS * (len(songs) - 1)
    full_total = total_secs + gap_total
    h, m = divmod(int(full_total), 3600)
    m, s = divmod(m, 60)
    summary = (
        f"총 {len(songs)}곡 | "
        f"음악 총 길이: {int(total_secs//60)}분 {int(total_secs%60):02d}초 | "
        f"영상 총 길이 (간격 포함): {h}시간 {m}분 {s:02d}초"
    )
    p = doc.add_paragraph(summary)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 빈 줄

    # 노래 목록 생성 (예: 00:00  01. Rain on the Window)
    current_time_secs = 0.0
    for i, song in enumerate(songs):
        h, m = divmod(int(current_time_secs), 3600)
        m, s = divmod(m, 60)
        if h > 0:
            start_time_str = f"{h}:{m:02d}:{s:02d}"
        else:
            start_time_str = f"{m:02d}:{s:02d}"

        line_text = f"{start_time_str}  {i + 1:02d}. {song['title']}"
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(2)
        p_line.paragraph_format.line_spacing = 1.15
        
        run = p_line.add_run(line_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(11)

        # 다음 노래의 시작 시간 계산 (재생 시간 + 간격)
        current_time_secs += song['duration_secs'] + GAP_SECONDS

    try:
        doc.save(output_path)
        print(f"📄 노래 목록 문서 저장: {output_path}\n")
    except PermissionError:
        print(f"⚠️  {output_path} 파일이 다른 프로그램(예: Word)에서 열려 있어 저장할 수 없습니다.")
        # 대안 파일명으로 저장 시도
        base_path = Path(output_path)
        saved_alt = False
        for i in range(1, 100):
            alt_path = base_path.with_name(f"{base_path.stem}_{i}{base_path.suffix}")
            try:
                doc.save(str(alt_path))
                print(f"📄 노래 목록 문서가 다른 이름으로 저장되었습니다: {alt_path}\n")
                saved_alt = True
                break
            except PermissionError:
                continue
        if not saved_alt:
            print("❌ 노래 목록 문서를 저장할 수 없었습니다.\n")


# ─────────────────────────────────────────────
#  메인
# ─────────────────────────────────────────────

def main():
    print("=" * 55)
    print("  CapCut 음악 편집 자동화 (pycapcut)")
    print("=" * 55)

    check_dependencies()

    songs  = get_songs(MUSIC_FOLDER)
    songs  = sort_and_shuffle_songs(songs)
    images = get_background_media(IMAGES_FOLDER)

    # 노래 목록 문서 생성
    docx_output = str(Path(MUSIC_FOLDER) / "노래목록.docx")
    generate_song_list_docx(songs, docx_output)

    # CapCut 초안 생성
    build_capcut_draft(songs, images)

    print("=" * 55)
    print("  완료! 다음 단계:")
    print("  1. CapCut을 열거나 재시작하세요")
    print(f"  2. 초안 목록에서 '{DRAFT_NAME}'을 찾아 여세요")
    print(f"  3. 노래 목록: {docx_output}")
    print("=" * 55)


if __name__ == "__main__":
    main()
